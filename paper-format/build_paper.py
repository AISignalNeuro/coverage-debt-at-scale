# -*- coding: utf-8 -*-
"""
Coverage Debt at Scale - rebuild the paper on the RA-L Word template
====================================================================
The submitted .docx was never built on an IEEE template: its styles.xml held
only Word defaults, 115 of 121 paragraphs were direct-formatted, and three of
its four sectPr blocks were A4. This rebuilds the same text on
ieeeconf_letter.dot without touching a single word.

Method:
  base.docx is Documents.Add(ieeeconf_letter.dot) -> it already carries the
  template's styles.xml, numbering.xml, settings.xml, header1.xml and sectPr.
  We empty its body and pour the content back into the template's OWN named
  styles. Nothing about page layout is hard-coded here.

  Text is lifted run-by-run out of the source .docx, never retyped.

Input:
  base.docx  (run make_base.ps1 first)
  src.docx   (the returned submission)
  figures/

Output:
  Coverage_Debt_at_Scale_ieeeconf.docx

Env: Python 3.12, Windows. Needs python-docx + lxml. Word only for make_base.
"""

import os
import re
import copy
import shutil

import docx
from docx.oxml.ns import qn, nsmap
from docx.shared import Inches
from lxml import etree

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "src.docx")
BASE = os.path.join(HERE, "base.docx")
OUT = os.path.join(HERE, "Coverage_Debt_at_Scale_ieeeconf.docx")
FIG = os.path.join(HERE, "figures")

W = nsmap["w"]
A = "http://schemas.openxmlformats.org/drawingml/2006/main"
WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
M = "http://schemas.openxmlformats.org/officeDocument/2006/math"

NS = ('xmlns:w="%s" xmlns:a="%s" xmlns:wps="%s" xmlns:wp="%s" xmlns:m="%s" '
      'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
      'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"') % (W, A, WPS, WP, M)


def X(xml):
    """parse a fragment with the namespaces bound"""
    tag = re.match(r"\s*<([\w:]+)", xml).group(1)
    return etree.fromstring(xml.replace("<%s" % tag, "<%s %s" % (tag, NS), 1))


def esc(t):
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# --- source side: read the returned document, verbatim ---
src = docx.Document(SRC)
src_paras = [p for p in src.element.body.iterchildren() if p.tag == qn("w:p")]
src_tbl = [t for t in src.element.body.iterchildren() if t.tag == qn("w:tbl")][0]


def runs_of(i):
    """[(text, bold, italic)] for source paragraph i, whitespace untouched"""
    p = docx.text.paragraph.Paragraph(src_paras[i], src)
    return [(r.text, bool(r.bold), bool(r.italic)) for r in p.runs if r.text]


def text_of(i):
    return "".join(t for t, _, _ in runs_of(i))


# --- start from the template ---
shutil.copyfile(BASE, OUT)
doc = docx.Document(OUT)
body = doc.element.body

sectPr = copy.deepcopy(body.find(qn("w:sectPr")))
for child in list(body):
    body.remove(child)

# register the three images once so we get valid r:embed ids + image parts
scratch = doc.add_paragraph()
rid, im_cx, im_cy = {}, {}, {}
for key, fn in [("f1", "fig1_heatmap_600dpi.png"),
                ("f2", "fig2_rarefaction.png"),
                ("f3", "fig3_debt.png")]:
    run = scratch.add_run()
    run.add_picture(os.path.join(FIG, fn))
    dr = run._r.findall(qn("w:drawing"))[-1]
    rid[key] = next(dr.iter("{%s}blip" % A)).get(qn("r:embed"))
    ext = next(dr.iter("{%s}extent" % WP))
    im_cx[key], im_cy[key] = int(ext.get("cx")), int(ext.get("cy"))
body.remove(scratch._p)

_n = [1000]


def uid():
    _n[0] += 1
    return _n[0]


def pic(key, width_in):
    """inline drawing scaled to width_in, true aspect ratio"""
    cx = int(Inches(width_in))
    cy = int(round(cx * im_cy[key] / im_cx[key]))
    i = uid()
    return ('<w:drawing><wp:inline distT="0" distB="0" distL="0" distR="0">'
            '<wp:extent cx="%d" cy="%d"/><wp:effectExtent l="0" t="0" r="0" b="0"/>'
            '<wp:docPr id="%d" name="Picture %d"/><wp:cNvGraphicFramePr>'
            '<a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr><a:graphic>'
            '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
            '<pic:pic><pic:nvPicPr><pic:cNvPr id="%d" name="Picture %d"/><pic:cNvPicPr/>'
            '</pic:nvPicPr><pic:blipFill><a:blip r:embed="%s"/><a:stretch><a:fillRect/>'
            '</a:stretch></pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/>'
            '<a:ext cx="%d" cy="%d"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/>'
            '</a:prstGeom></pic:spPr></pic:pic></a:graphicData></a:graphic>'
            '</wp:inline></w:drawing>' % (cx, cy, i, i, i, i, rid[key], cx, cy)), cy


# --- paragraph builders: everything gets a template pStyle ---
def runs_xml(parts):
    out = []
    for t, b, i in parts:
        rpr = ("<w:b/>" if b else "") + ("<w:i/>" if i else "")
        rpr = "<w:rPr>%s</w:rPr>" % rpr if rpr else ""
        out.append('<w:r>%s<w:t xml:space="preserve">%s</w:t></w:r>' % (rpr, esc(t)))
    return "".join(out)


def para(style, parts, ppr=""):
    if isinstance(parts, str):
        parts = [(parts, False, False)]
    return X('<w:p><w:pPr><w:pStyle w:val="%s"/>%s</w:pPr>%s</w:p>'
             % (style, ppr, runs_xml(parts)))


def add(el):
    body.append(el)
    return el


# --- OMML bits for the two numbered equations ---
MFONT = '<w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr>'


def mr(text, plain=False):
    """math run; plain=True -> upright (numbers, operators, delimiters)"""
    sty = '<m:rPr><m:sty m:val="p"/></m:rPr>' if plain else ""
    return "<m:r>%s%s<m:t>%s</m:t></m:r>" % (sty, MFONT, esc(text))


def msub(base, sub):
    return ('<m:sSub><m:sSubPr><m:ctrlPr>%s</m:ctrlPr></m:sSubPr><m:e>%s</m:e>'
            "<m:sub>%s</m:sub></m:sSub>" % (MFONT, base, sub))


# --- column-spanning float: image + caption in a borderless box, top of page ---
def float_figure(key, width_in, caption):
    drawing, cy = pic(key, width_in)
    box_cx, box_cy = int(Inches(7.0)), cy + int(Inches(0.75))
    i = uid()
    inner = ('<w:p><w:pPr><w:jc w:val="center"/><w:spacing w:after="60" w:line="240" '
             'w:lineRule="auto"/></w:pPr><w:r>%s</w:r></w:p>'
             '<w:p><w:pPr><w:pStyle w:val="FigureCaption0"/><w:jc w:val="both"/>'
             '<w:spacing w:after="0"/></w:pPr>%s</w:p>' % (drawing, runs_xml(caption)))
    return X('<w:p><w:pPr><w:pStyle w:val="BodyText"/><w:spacing w:after="0" w:line="240" '
             'w:lineRule="auto"/><w:ind w:firstLine="0"/></w:pPr><w:r><w:drawing>'
             '<wp:anchor distT="0" distB="182880" distL="114300" distR="114300" simplePos="0" '
             'relativeHeight="%d" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="0">'
             '<wp:simplePos x="0" y="0"/>'
             '<wp:positionH relativeFrom="margin"><wp:align>center</wp:align></wp:positionH>'
             '<wp:positionV relativeFrom="margin"><wp:align>top</wp:align></wp:positionV>'
             '<wp:extent cx="%d" cy="%d"/><wp:effectExtent l="0" t="0" r="0" b="0"/>'
             '<wp:wrapTopAndBottom/><wp:docPr id="%d" name="Figure box %d"/>'
             '<wp:cNvGraphicFramePr/><a:graphic><a:graphicData uri="%s"><wps:wsp>'
             '<wps:cNvSpPr txBox="1"/><wps:spPr><a:xfrm><a:off x="0" y="0"/>'
             '<a:ext cx="%d" cy="%d"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
             '<a:noFill/><a:ln><a:noFill/></a:ln></wps:spPr>'
             '<wps:txbx><w:txbxContent>%s</w:txbxContent></wps:txbx>'
             '<wps:bodyPr rot="0" vert="horz" wrap="square" lIns="0" tIns="0" rIns="0" '
             'bIns="0" anchor="t" anchorCtr="0"><a:spAutoFit/></wps:bodyPr>'
             "</wps:wsp></a:graphicData></a:graphic></wp:anchor></w:drawing></w:r></w:p>"
             % (251650000 + i, box_cx, box_cy, i, i, WPS, box_cx, box_cy, inner))


# --------------------------- compose ---------------------------
# Title and Authors are page-width frames in this template. It anchors them on a
# 1pt `Text` paragraph and overrides framePr on each -- skip that and the two
# frames collide and come out in the wrong order.
add(X('<w:p><w:pPr><w:pStyle w:val="Text"/><w:ind w:firstLine="0"/>'
      '<w:rPr><w:sz w:val="2"/><w:szCs w:val="18"/></w:rPr></w:pPr></w:p>'))
add(para("Title", text_of(1), '<w:framePr w:wrap="notBeside"/>'))
add(para("Authors", text_of(2), '<w:framePr w:wrap="notBeside" w:x="1614"/>'))


def labelled(i):
    """IEEE sets the leading Abstract / Index Terms label in italic"""
    parts = runs_of(i)
    return [(parts[0][0], False, True)] + [(t, False, False) for t, _, _ in parts[1:]]


add(para("Abstract", labelled(5), '<w:spacing w:before="0"/>'))
add(para("IndexTerms", labelled(6)))

# Heading1/Heading2 auto-number in this template, so the literal numeral has to
# go or it doubles. \s (not " ") because the source separates it with an em space.
NUM = re.compile(r"^(?:[IVX]+|[A-D])\.\s+")


def heading(i, level):
    text = text_of(i)
    stripped = NUM.sub("", text)
    assert stripped != text, "heading numeral not stripped: %r" % text[:40]
    return para("Heading%d" % level, stripped)


# I. Introduction
add(heading(7, 1))
for i in (8, 9, 10, 11, 12):
    add(para("BodyText", runs_of(i)))

# II. Related Work
add(heading(13, 1))
for i in (14, 15, 16, 17):
    add(para("BodyText", runs_of(i)))

# III. Method
add(heading(18, 1))
add(heading(19, 2))
add(para("BodyText", runs_of(20)))
add(heading(21, 2))
add(para("BodyText", runs_of(22)))
add(heading(23, 2))
add(para("BodyText", runs_of(24)))

EQ_TABS = ('<w:tabs><w:tab w:val="center" w:pos="2405"/>'
           '<w:tab w:val="right" w:pos="4810"/></w:tabs>')

eq1 = ("<m:oMath>" + mr("b") + mr("(", True) + mr("s") + mr(")", True)
       + mr(" = (1/|", True) + mr("E") + mr("|) ", True)
       + msub(mr("Σ", True), mr("e"))
       + mr("c") + mr("(", True) + mr("e") + mr(",", True) + mr("s") + mr("),", True)
       + mr("  ", True)
       + mr("d") + mr("(", True) + mr("s") + mr(")", True) + mr(" = 1 − ", True)
       + mr("b") + mr("(", True) + mr("s") + mr("),", True) + "</m:oMath>")
add(X('<w:p><w:pPr><w:pStyle w:val="Equation"/>%s</w:pPr><w:r><w:tab/></w:r>%s'
      "<w:r><w:tab/><w:t>(1)</w:t></w:r></w:p>" % (EQ_TABS, eq1)))

add(para("BodyText", runs_of(26), '<w:ind w:firstLine="0"/>'))

eq2 = ("<m:oMath>" + mr("SCD", True) + mr(" = (1/|", True) + mr("S") + mr("|) ", True)
       + msub(mr("Σ", True), mr("s"))
       + mr("d") + mr("(", True) + mr("s") + mr(").", True) + "</m:oMath>")
add(X('<w:p><w:pPr><w:pStyle w:val="Equation"/>%s</w:pPr><w:r><w:tab/></w:r>%s'
      "<w:r><w:tab/><w:t>(2)</w:t></w:r></w:p>" % (EQ_TABS, eq2)))

add(para("BodyText", runs_of(28), '<w:ind w:firstLine="0"/>'))
add(heading(29, 2))
add(para("BodyText", runs_of(30)))

# IV. Results -- Fig. 2 is cited in IV-A, anchor its float just ahead of it
add(float_figure("f2", 6.9, runs_of(42)))
add(heading(31, 1))
add(heading(32, 2))
add(para("BodyText", runs_of(33)))

add(float_figure("f1", 6.9, runs_of(40)))

# Fig. 3 goes single column at the head of IV-B, which cites it. Putting it after
# the citing paragraph instead leaves ~2in of white at the foot of the left column.
p3, _ = pic("f3", 3.25)
add(X('<w:p><w:pPr><w:pStyle w:val="BodyText"/><w:jc w:val="center"/>'
      '<w:ind w:firstLine="0"/><w:spacing w:before="0" w:after="60"/><w:keepNext/>'
      "</w:pPr><w:r>%s</w:r></w:p>" % p3))
add(para("FigureCaption0", runs_of(45), "<w:keepNext/>"))

add(heading(34, 2))
add(para("BodyText", runs_of(35)))
add(heading(46, 2))
add(para("BodyText", runs_of(47)))
add(heading(48, 2))
add(para("BodyText", runs_of(49)))

# V. Discussion
add(heading(50, 1))
for i in (51, 52):
    add(para("BodyText", runs_of(i)))

# VI. Conclusion
add(heading(53, 1))
add(para("BodyText", runs_of(54)))
add(para("BodyText", runs_of(55)))

# References: heading unnumbered, then the auto-numbered [n] list
add(para("Heading1", "References",
         '<w:numPr><w:ilvl w:val="0"/><w:numId w:val="0"/></w:numPr>'))
# same story: References auto-numbers [n], so strip the literal one
REF = re.compile(r"^\[\d+\]\s+")
for i in range(57, 97):
    text = text_of(i)
    stripped = REF.sub("", text)
    assert stripped != text, "reference numeral not stripped: %r" % text[:40]
    add(para("References", stripped))

# --- Table I: IEEE rules (horizontal only), floating, spans both columns ---
rows_src = [[c.text for c in r.cells] for r in docx.table.Table(src_tbl, src).rows]

COLW = [1500, 2560, 2560, 3220]          # twips; 9840 total = 6.83in
RULE = '<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
NONE = ('<w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/>'
        '<w:right w:val="nil"/>')


def cell(text, style, col=0, span=None, borders="", jc=None):
    tcpr = '<w:tcW w:w="%d" w:type="dxa"/>' % (sum(COLW) if span else COLW[col])
    if span:
        tcpr += '<w:gridSpan w:val="%d"/>' % span
    if borders:
        tcpr += "<w:tcBorders>%s</w:tcBorders>" % borders
    tcpr += '<w:vAlign w:val="center"/>'
    if isinstance(text, str):
        text = [text]
    ps = "".join('<w:p><w:pPr><w:pStyle w:val="%s"/>%s<w:spacing w:before="20" '
                 'w:after="20" w:line="240" w:lineRule="auto"/></w:pPr>'
                 '<w:r><w:t xml:space="preserve">%s</w:t></w:r></w:p>'
                 % (style, '<w:jc w:val="%s"/>' % jc if jc else "", esc(t)) for t in text)
    return "<w:tc><w:tcPr>%s</w:tcPr>%s</w:tc>" % (tcpr, ps)


cap = text_of(37)
rows = ['<w:tr><w:trPr><w:cantSplit/></w:trPr>%s</w:tr>'
        % cell(["TABLE I", cap.split("TABLE I", 1)[1].strip()], "TableTitle",
               span=4, borders=NONE)]
rows.append('<w:tr><w:trPr><w:cantSplit/><w:tblHeader/></w:trPr>%s</w:tr>' % "".join(
    cell(t, "tablecolhead", col=ci,
         borders=RULE + '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>')
    for ci, t in enumerate(rows_src[0])))
for ri, row in enumerate(rows_src[1:], start=1):
    b = ('<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
         if ri == len(rows_src) - 1 else "")
    rows.append('<w:tr><w:trPr><w:cantSplit/></w:trPr>%s</w:tr>'
                % "".join(cell(t, "tablecopy", col=ci, borders=b,
                               jc="left" if ci else "center")
                          for ci, t in enumerate(row)))

table_xml = ('<w:tbl><w:tblPr><w:tblpPr w:leftFromText="180" w:rightFromText="180" '
             'w:vertAnchor="margin" w:horzAnchor="margin" w:tblpXSpec="center" '
             'w:tblpYSpec="top"/><w:tblW w:w="%d" w:type="dxa"/><w:jc w:val="center"/>'
             '<w:tblBorders><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/>'
             '<w:right w:val="nil"/><w:insideH w:val="nil"/><w:insideV w:val="nil"/>'
             '</w:tblBorders><w:tblCellMar><w:top w:w="20" w:type="dxa"/>'
             '<w:left w:w="72" w:type="dxa"/><w:bottom w:w="20" w:type="dxa"/>'
             '<w:right w:w="72" w:type="dxa"/></w:tblCellMar><w:tblLook w:val="0000"/>'
             "</w:tblPr><w:tblGrid>%s</w:tblGrid>%s</w:tbl>"
             % (sum(COLW), "".join('<w:gridCol w:w="%d"/>' % w for w in COLW), "".join(rows)))

# Anchor it where it is first cited (Sec. II). Match on collapsed whitespace --
# the source writes "Table I" with a narrow no-break space, not a plain one.
ANCHOR = "Table I contrasts the three views."


def flat(el):
    return re.sub(r"\s+", " ", "".join(t.text or "" for t in el.iter(qn("w:t"))))


for p in body.findall(qn("w:p")):
    if ANCHOR in flat(p):
        p.addprevious(X(table_xml))
        break
else:
    raise SystemExit("could not find the Table I anchor paragraph")

body.append(sectPr)

# --- scrub document properties ---
# base.docx inherits whatever the local Office install and the template left
# behind: this machine's Word stamps dc:creator "RePack by Diakov" and the
# template carries subject "IEEE Transactions on Magnetics". Both ride into the
# exported PDF, and author metadata is exactly what de-anonymises a blind
# submission. Blank rather than the real name, since the byline is still open.
cp = doc.core_properties
for field in ("author", "last_modified_by", "title", "subject", "keywords",
              "comments", "category", "content_status"):
    setattr(cp, field, "")
cp.revision = 1

doc.save(OUT)
print("wrote", OUT)
